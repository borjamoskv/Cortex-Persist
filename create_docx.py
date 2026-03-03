from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


def set_cell_background(cell, color_hex):
    # Set cell background color
    shading_elm = OxmlElement("w:shd")
    shading_elm.set(qn("w:fill"), color_hex)
    cell._tc.get_or_add_tcPr().append(shading_elm)


def create_document():
    doc = Document()

    # Configurar márgenes
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Estilos
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Inter"
    font.size = Pt(11)

    # Portada
    doc.add_heading("", level=0)
    title = doc.add_heading("CORTEX-PERSIST", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph(
        "Infraestructura de Confianza Descentralizada para Agentes de IA Autónomos"
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(16)
    subtitle.runs[0].font.color.rgb = RGBColor(46, 80, 144)  # YInMn Blue

    doc.add_paragraph("\n\n\n\n")
    author = doc.add_paragraph(
        "Análisis Técnico Exhaustivo\nAutor: Borja Fernández Angulo\nCORTEX System v4.0"
    )
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # Secciones principales
    section_titles = [
        "1. El Imperativo Estratégico: La Crisis de Confianza",
        "2. Contexto Regulatorio: Cumplimiento por Diseño y EU AI Act",
        "3. El Motor de Confianza CORTEX (arquitectura de 5 capas)",
        "4. Sistema de Memoria Cognitiva Tripartita (L1/L2/L3)",
        "5. Criptografía y Verificación (SHA-256, Merkle Trees, Sandbox AST)",
        "6. Consenso Multi-Agente (WBFT)",
        "7. Interoperabilidad y Despliegue Multiplataforma",
        "8. Metrología del Código Base (45,500 LOC, 444 módulos)",
        "9. Conclusiones Estratégicas",
    ]

    doc.add_heading("Tabla de Contenidos", level=1)
    for _index, sec in enumerate(section_titles):
        p = doc.add_paragraph(sec)
        p.style = doc.styles["List Number"]

    doc.add_page_break()

    # Añadir contenido por sección con 6 tablas en total

    # Sección 1
    doc.add_heading(section_titles[0], level=1)
    doc.add_paragraph(
        "La adopción de agentes de inteligencia artificial autónomos en entornos empresariales críticos se ve frenada por una crisis fundamental: la falta de confianza. Las organizaciones requieren garantías irrefutables de que los agentes operan dentro de límites predefinidos, que sus acciones son trazables y que la infraestructura subyacente puede resistir manipulaciones maliciosas o derivas estocásticas. CORTEX-Persist emerge como la solución arquitectónica a esta anomalía estructurada."
    )

    # Sección 2
    doc.add_heading(section_titles[1], level=1)
    doc.add_paragraph(
        "CORTEX implementa un modelo de 'Cumplimiento por Diseño', asegurando que la arquitectura base satisface nativamente los requisitos regulatorios globales, en particular la EU AI Act."
    )
    doc.add_heading("Tabla 1: Cumplimiento Normativo (EU AI Act)", level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    table.columns[0].width = Inches(2.5)
    table.columns[1].width = Inches(4.0)

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Requisito EU AI Act"
    hdr_cells[1].text = "Implementación en CORTEX"
    set_cell_background(hdr_cells[0], "2E5090")  # Navy blue header
    set_cell_background(hdr_cells[1], "2E5090")
    hdr_cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    hdr_cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    row_data = [
        (
            "Transparencia y Trazabilidad",
            "Log inmutable criptográfico de decisiones y acciones de los agentes (Chain of Thought Ledger).",
        ),
        (
            "Supervisión Humana",
            'Protocolos "Human-in-the-loop" obligatorios en operaciones críticas (C-Level Sandbox).',
        ),
        (
            "Gestión de Riesgos",
            "Evaluación estocástica y sandboxing estricto de ejecución de código (AST validation).",
        ),
    ]
    for req, imp in row_data:
        row_cells = table.add_row().cells
        row_cells[0].text = req
        row_cells[1].text = imp

    # Sección 3
    doc.add_heading(section_titles[2], level=1)
    doc.add_paragraph(
        "El núcleo de CORTEX es su motor de confianza, segmentado en cinco bi-capas estructurales para garantizar aislamiento absoluto y comunicación segura (Zero-Trust)."
    )
    doc.add_heading("Tabla 2: Interfaces de las Capas del Motor CORTEX", level=2)
    table2 = doc.add_table(rows=1, cols=2)
    table2.style = "Table Grid"
    table2.autofit = False
    table2.columns[0].width = Inches(1.5)
    table2.columns[1].width = Inches(5.0)

    hdr_cells2 = table2.rows[0].cells
    hdr_cells2[0].text = "Capa"
    hdr_cells2[1].text = "Responsabilidad e Interfaz"
    set_cell_background(hdr_cells2[0], "2E5090")
    set_cell_background(hdr_cells2[1], "2E5090")
    hdr_cells2[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    hdr_cells2[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    layers = [
        ("L0 - Base", "Sistema subyacente y sistema de archivos cifrado (AES-256)."),
        ("L1 - Engine", "Motor de ejecución determinista y Sandbox AST."),
        ("L2 - Memory", "Memoria transaccional L1/L2/L3 (Redis/SQLite)."),
        ("L3 - Agent", "Máquina de estados finitos del Agente y WBFT."),
        ("L4 - Network", "Protocolos M2M cifrados y RPC sobre TLS 1.3."),
    ]
    for layer, desc in layers:
        row_cells = table2.add_row().cells
        row_cells[0].text = layer
        row_cells[1].text = desc

    # Sección 4
    doc.add_heading(section_titles[3], level=1)
    doc.add_paragraph(
        "El sistema de memoria simula la cognición mamífera avanzada, permitiendo a los agentes razonar temporalmente sin colapso de contexto."
    )
    doc.add_heading("Tabla 3: Niveles de Memoria Cognitiva", level=2)
    table3 = doc.add_table(rows=1, cols=3)
    table3.style = "Table Grid"
    table3.autofit = False
    for col in table3.columns:
        col.width = Inches(2.1)

    hdr_cells3 = table3.rows[0].cells
    hdr_cells3[0].text = "Nivel Cognitivo"
    hdr_cells3[1].text = "Tecnología"
    hdr_cells3[2].text = "Persistencia / TTL"
    for cell in hdr_cells3:
        set_cell_background(cell, "2E5090")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    mem_levels = [
        ("L1: Working Memory", "Redis (In-Memory)", "Efímera (Volátil, sub-milisegundo)."),
        (
            "L2: Episodic Memory",
            "SQLiteDB (Transaccional)",
            "Sesión / Proyecto (Indices semánticos).",
        ),
        ("L3: Semantic Core", "Vector DB", "Permanente (Conocimiento ontológico)."),
    ]
    for lvl, tech, ttl in mem_levels:
        row_cells = table3.add_row().cells
        row_cells[0].text = lvl
        row_cells[1].text = tech
        row_cells[2].text = ttl

    # Sección 5
    doc.add_heading(section_titles[4], level=1)
    doc.add_paragraph(
        "La integridad de los datos en CORTEX está garantizada a través de primitivas criptográficas robustas y un entorno de validación abstracto (AST). Cada modificación de estado genera un hash que se entrelaza en un Árbol de Merkle."
    )

    # Sección 6
    doc.add_heading(section_titles[5], level=1)
    doc.add_paragraph(
        "La comunicación en enjambre requiere tolerancia a fallos. CORTEX implementa una variante de Tolerancia a Faltas Bizantinas Ponderada (WBFT)."
    )
    doc.add_heading("Tabla 4: Frameworks de Consenso", level=2)
    table4 = doc.add_table(rows=1, cols=2)
    table4.style = "Table Grid"
    table4.autofit = False
    table4.columns[0].width = Inches(2.0)
    table4.columns[1].width = Inches(4.5)

    hdr_cells4 = table4.rows[0].cells
    hdr_cells4[0].text = "Mecanismo"
    hdr_cells4[1].text = "Implementación CORTEX-WBFT"
    for cell in hdr_cells4:
        set_cell_background(cell, "2E5090")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    consensus = [
        (
            "Leader Election",
            "Rotación determinista basada en reputación histórica (proof-of-accuracy).",
        ),
        ("Quorum", "Super-mayoría 2/3 para decisiones de Nivel 3 (Modificaciones de código)."),
        ("Slashing", "Penalización de tokens operacionales para agentes divergentes o maliciosos."),
    ]
    for mech, imp in consensus:
        row_cells = table4.add_row().cells
        row_cells[0].text = mech
        row_cells[1].text = imp

    # Sección 7
    doc.add_heading(section_titles[6], level=1)
    doc.add_paragraph("La arquitectura de CORTEX soporta portabilidad nativa y resiliencia de red.")
    doc.add_heading("Tabla 5: Estrategias de Despliegue", level=2)
    table5 = doc.add_table(rows=1, cols=2)
    table5.style = "Table Grid"
    table5.autofit = False
    table5.columns[0].width = Inches(2.0)
    table5.columns[1].width = Inches(4.5)

    hdr_cells5 = table5.rows[0].cells
    hdr_cells5[0].text = "Entorno"
    hdr_cells5[1].text = "Técnica de Aislamiento"
    for cell in hdr_cells5:
        set_cell_background(cell, "2E5090")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    deploy = [
        ("Bare-Metal (Desktop)", "Local Daemon con privilegios limitados y chroot (NotchLive)."),
        ("Nube (K8s/Docker)", "Microservicios L1/L2/L3 orquestados mediante CORTEX Nexus."),
        ("Edge/IoT", "CORTEX-Lite (Rust footprint de baja latencia)."),
    ]
    for env, tech in deploy:
        row_cells = table5.add_row().cells
        row_cells[0].text = env
        row_cells[1].text = tech

    # Sección 8
    doc.add_heading(section_titles[7], level=1)
    doc.add_paragraph(
        "Para mantener la robustez sin caer en el gigantismo de software, CORTEX restringe su código base enfocándose en la densidad y el determinismo."
    )
    doc.add_heading("Tabla 6: Métricas Base de Código", level=2)
    table6 = doc.add_table(rows=1, cols=2)
    table6.style = "Table Grid"
    table6.autofit = False
    table6.columns[0].width = Inches(2.5)
    table6.columns[1].width = Inches(4.0)

    hdr_cells6 = table6.rows[0].cells
    hdr_cells6[0].text = "Métrica (CORTEX v4)"
    hdr_cells6[1].text = "Valor"
    for cell in hdr_cells6:
        set_cell_background(cell, "2E5090")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    metrics = [
        ("Líneas de Código (LOC)", "45,500 (Densidad optimizada)"),
        ("Módulos Estructurales", "444 componentes aislados"),
        ("Cobertura de Tests", "> 94% (Unit, E2E, Chaos)"),
        ("Latencia Promedio L1", "< 15ms (P99)"),
    ]
    for met, val in metrics:
        row_cells = table6.add_row().cells
        row_cells[0].text = met
        row_cells[1].text = val

    # Sección 9
    doc.add_heading(section_titles[8], level=1)
    doc.add_paragraph(
        "La adopción de CORTEX-Persist marca la transición de simples bots conversacionales de IA a flotas de agentes autónomos económicamente viables y legalmente defendibles. La asimetría entrópica resuelta por su arquitectura L1 a L3 posiciona a las organizaciones como inmunes a las ineficiencias de enjambres no estructurados y riesgos regulatorios masivos."
    )

    doc.add_page_break()

    # Contraportada
    doc.add_heading("Contraportada", level=1)
    doc.add_paragraph(
        "Proyecto CORTEX-PERSIST\n"
        "Arquitectura Soberana MOSKV-1\n"
        "Licencia Privada (Strictly Confidential)\n\n"
        "© 2026 Borja Fernández Angulo"
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER

    try:
        doc.save("Cortex_Persist_Analisis.docx")
        print("Documento guardado: Cortex_Persist_Analisis.docx")
    except Exception as e:
        print(f"Error saving document: {e}")


if __name__ == "__main__":
    create_document()
